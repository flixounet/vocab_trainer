# streamlit_app.py — stabile Version
import streamlit as st
from io import BytesIO
import json, random, unicodedata

# --------- optionaler Word-Import ----------
try:
    from docx import Document
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False

# --------- Hilfsfunktionen ----------
def normalize(s: str) -> str:
    s = s.strip().lower()
    s = " ".join(s.split())
    s = unicodedata.normalize("NFKD", s)
    s = "".join(ch for ch in s if not unicodedata.combining(ch))
    return s

DEFAULT_STORE = {
    "collections": [{
        "name": "Evolution_und_Steinzeit",
        "items": [
            {"de": "die Urgeschichte", "fr": "la Préhistoire"},
            {"de": "die Frühgeschichte", "fr": "la Protohistoire"},
            {"de": "die Altsteinzeit (2,5 Mio.-9500 v. Chr.)", "fr": "le Paléolithique"},
            {"de": "die Jungsteinzeit (9500 v. Chr.-2200 v. Chr.)", "fr": "le Néolithique"},
            {"de": "der Archäologe", "fr": "l'archéologue"},
            {"de": "die Höhlenmalerei", "fr": "la peinture pariétale"},
            {"de": "der Nomade, die Nomadin", "fr": "un/une nomade"},
            {"de": "roden, urbar machen", "fr": "défricher"},
            {"de": "der/die Sesshafte", "fr": "le/la sédentaire"},
            {"de": "sesshaft werden", "fr": "devenir sédentaire"},
            {"de": "der Tauschhandel", "fr": "le troc"},
            {"de": "der Jäger und Sammler", "fr": "le chasseur-cueilleur"},
            {"de": "der Faustkeil", "fr": "le biface en silex"},
            {"de": "das Haustier", "fr": "l'animal domestique"},
        ]
    }]
}

def import_docx(file_bytes, filename):
    if not DOCX_AVAILABLE:
        st.error("python-docx ist nicht installiert – bitte in requirements.txt ergänzen.")
        return None
    doc = Document(BytesIO(file_bytes))
    items = []

    # Tabellen (2 Spalten: DE | FR)
    for tbl in doc.tables:
        for r_i, row in enumerate(tbl.rows):
            cells = [c.text.strip() for c in row.cells]
            if len(cells) >= 2:
                de, fr = cells[0], cells[1]
                if not de or not fr:
                    continue
                if r_i == 0 and ("de" in de.lower() and "fr" in fr.lower()):
                    continue
                items.append({"de": de, "fr": fr})

    # Absätze "de ; fr"
    for p in doc.paragraphs:
        t = p.text.strip()
        if ";" in t:
            parts = [s.strip() for s in t.split(";")]
            if len(parts) >= 2 and parts[0] and parts[1]:
                items.append({"de": parts[0], "fr": parts[1]})

    # Dedupe
    seen = set(); uniq = []
    for it in items:
        key = (normalize(it["de"]), normalize(it["fr"]))
        if key in seen: 
            continue
        seen.add(key); uniq.append(it)
    return {"name": filename.rsplit(".",1)[0], "items": uniq}

# --------- Streamlit Setup ----------
st.set_page_config(page_title="VocabQuiz DE↔FR", page_icon="🔤", layout="wide")
st.title("VocabQuiz – Deutsch ↔ Français")

# Persistent Store
st.session_state.setdefault("store", DEFAULT_STORE)

# Quiz-State: alles zentral in session_state (keine Kopien!)
st.session_state.setdefault("quiz_active", False)
st.session_state.setdefault("quiz_direction", "DE→FR")
st.session_state.setdefault("quiz_mode", "Multiple Choice")
st.session_state.setdefault("quiz_pool", [])
st.session_state.setdefault("quiz_order", [])   # Liste von Indexen in quiz_pool
st.session_state.setdefault("quiz_i", 0)        # aktuelle Position in quiz_order
st.session_state.setdefault("quiz_score", 0)
st.session_state.setdefault("quiz_history", []) # (Frage, AntwortUser, OK, Richtig)

def start_quiz(pool, direction, mode, n_q):
    st.session_state.quiz_direction = direction
    st.session_state.quiz_mode = mode
    st.session_state.quiz_pool = pool[:]  # Kopie
    # Reihenfolge als Indexliste — garantiert deterministisches Weiterblättern
    order = list(range(len(pool)))
    random.shuffle(order)
    st.session_state.quiz_order = order[:min(n_q, len(pool))]
    st.session_state.quiz_i = 0
    st.session_state.quiz_score = 0
    st.session_state.quiz_history = []
    st.session_state.quiz_active = True

def current_item():
    i = st.session_state.quiz_i
    pool = st.session_state.quiz_pool
    order = st.session_state.quiz_order
    if i >= len(order):
        return None
    return pool[order[i]]

def advance(item_question, user_answer, correct_answer, correct_bool):
    # Ergebnis speichern
    st.session_state.quiz_history.append(
        (item_question, user_answer, correct_bool, correct_answer)
    )
    if correct_bool:
        st.session_state.quiz_score += 1
    # zur nächsten Frage
    st.session_state.quiz_i += 1

# --------- Tabs ----------
tab_quiz, tab_manage = st.tabs(["🎯 Quiz", "📚 Sammlungen & Import"])

# ====== Sammlungen & Import ======
with tab_manage:
    st.subheader("Bestehende Sammlungen")
    for c in st.session_state.store.get("collections", []):
        st.markdown(f"- **{c.get('name')}** – {len(c.get('items', []))} Einträge")
    st.divider()

    st.subheader("Import aus Word (.docx)")
    up = st.file_uploader("Word-Datei hochladen", type=["docx"])
    if up is not None:
        data = import_docx(up.read(), up.name)
        if data and len(data["items"]) > 0:
            if st.button(f"Sammlung '{data['name']}' importieren ({len(data['items'])} Einträge)"):
                cols = st.session_state.store["collections"]
                idx = next((i for i, c in enumerate(cols) if c.get("name") == data["name"]), None)
                if idx is not None:
                    cols[idx] = data
                else:
                    cols.append(data)
                st.success(f"Importiert: {len(data['items'])} Einträge in '{data['name']}'")

    st.divider()
    st.subheader("Datenbank exportieren")
    js = json.dumps(st.session_state.store, ensure_ascii=False, indent=2).encode("utf-8")
    st.download_button("vocab_store.json herunterladen", js, file_name="vocab_store.json", mime="application/json")

# ====== Quiz ======
with tab_quiz:
    # Konfiguration
    options_collections = ["(alle)"] + [c.get("name","?") for c in st.session_state.store.get("collections", [])]
    coll = st.selectbox("Sammlung", options_collections, index=0, key="cfg_coll")
    direction = st.radio("Richtung", ["DE→FR", "FR→DE"], horizontal=True, key="cfg_dir")
    mode = st.radio("Quiztyp", ["Multiple Choice", "Freitext"], horizontal=True, key="cfg_mode")
    n_q = st.slider("Anzahl Fragen", 5, 50, 10, key="cfg_n")

    # Pool zusammenstellen
    pool = []
    for c in st.session_state.store.get("collections", []):
        if coll != "(alle)" and c.get("name") != coll:
            continue
        pool.extend(c.get("items", []))

    # Start-Button
    if st.button("Quiz starten", type="primary", key="start_btn") and len(pool) >= 4:
        start_quiz(pool, direction, mode, n_q)
        st.rerun()

    # Kein aktives Quiz?
    if not st.session_state.quiz_active:
        st.info("Konfiguration wählen und **Quiz starten**.")
        st.stop()

    # Aktuelle Frage
    item = current_item()
    total = len(st.session_state.quiz_order)

    if item is None:
        # Auswertung
        score = st.session_state.quiz_score
        st.success(f"Fertig! Punktzahl: {score}/{total}  ({round(100*score/max(1,total))}%)")
        st.dataframe(
            [{"Frage":h[0], "Ihre Antwort":h[1], "Korrekt":"Ja" if h[2] else "Nein", "Richtig":h[3]}
             for h in st.session_state.quiz_history],
            use_container_width=True
        )
        col1, col2 = st.columns(2)
        if col1.button("Neues Quiz", use_container_width=True, key="new_quiz"):
            st.session_state.quiz_active = False
            st.rerun()
        if col2.button("Nochmal gleiche Auswahl", use_container_width=True, key="repeat_quiz"):
            # Gleiche Einstellungen, aber neue Reihenfolge
            start_quiz(
                st.session_state.quiz_pool,
                st.session_state.quiz_direction,
                st.session_state.quiz_mode,
                n_q=len(st.session_state.quiz_order)
            )
            st.rerun()
        st.stop()

    # Anzeige der Frage
    qdir = st.session_state.quiz_direction
    question = item["de"] if qdir == "DE→FR" else item["fr"]
    correct = item["fr"] if qdir == "DE→FR" else item["de"]

    st.info(f"Frage {st.session_state.quiz_i+1}/{total}  •  Punktzahl: {st.session_state.quiz_score}")
    st.write(f"**Übersetze:** {question}")

    # Antwortmöglichkeiten bauen (aus dem ganzen Pool)
    if qdir == "DE→FR":
        all_answers = list({ x["fr"] for x in st.session_state.quiz_pool })
    else:
        all_answers = list({ x["de"] for x in st.session_state.quiz_pool })
    distractors = [a for a in all_answers if normalize(a) != normalize(correct)]
    random.shuffle(distractors)

    with st.form(key=f"form_{st.session_state.quiz_i}", clear_on_submit=True):
        if st.session_state.quiz_mode == "Multiple Choice":
            opts = [correct] + distractors[:3]
            random.shuffle(opts)
            user_answer = st.radio("Wähle die richtige Übersetzung:", opts, index=None, key=f"radio_{st.session_state.quiz_i}")
        else:
            user_answer = st.text_input("Antwort eingeben", key=f"text_{st.session_state.quiz_i}")

        submitted = st.form_submit_button("Weiter", use_container_width=True)

    if submitted:
        if (st.session_state.quiz_mode == "Multiple Choice" and user_answer is None) or \
           (st.session_state.quiz_mode == "Freitext" and not user_answer.strip()):
            st.warning("Bitte eine Antwort eingeben/auswählen.")
        else:
            ok = normalize(user_answer) == normalize(correct)
            advance(question, user_answer, correct, ok)
            st.rerun()
