# streamlit_app.py — brand-new stable build
import streamlit as st
from io import BytesIO
import json, random, unicodedata

try:
    from docx import Document
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False

APP_TITLE = "VocabQuiz – Deutsch ↔ Français"

def normalize(s: str) -> str:
    s = s.strip().lower()
    s = " ".join(s.split())
    s = unicodedata.normalize("NFKD", s)
    s = "".join(ch for ch in s if not unicodedata.combining(ch))
    return s

DEFAULT_STORE = {
    "collections": [{
        "name": "Evolution_und_Steinzeit",
        "items": [
            {"de": "die Urgeschichte", "fr": "la Préhistoire"},
            {"de": "die Frühgeschichte", "fr": "la Protohistoire"},
            {"de": "die Altsteinzeit (2,5 Mio.-9500 v. Chr.)", "fr": "le Paléolithique"},
            {"de": "die Jungsteinzeit (9500 v. Chr.-2200 v. Chr.)", "fr": "le Néolithique"},
            {"de": "der Archäologe", "fr": "l'archéologue"},
            {"de": "die Höhlenmalerei", "fr": "la peinture pariétale"},
            {"de": "der Nomade, die Nomadin", "fr": "un/une nomade"},
            {"de": "roden, urbar machen", "fr": "défricher"},
            {"de": "der/die Sesshafte", "fr": "le/la sédentaire"},
            {"de": "sesshaft werden", "fr": "devenir sédentaire"},
            {"de": "der Tauschhandel", "fr": "le troc"},
            {"de": "der Jäger und Sammler", "fr": "le chasseur-cueilleur"},
            {"de": "der Faustkeil", "fr": "le biface en silex"},
            {"de": "das Haustier", "fr": "l'animal domestique"},
        ]
    }]
}

def import_docx(file_bytes, filename):
    if not DOCX_AVAILABLE:
        st.error("python-docx ist nicht installiert – bitte in requirements.txt ergänzen.")
        return None
    doc = Document(BytesIO(file_bytes))
    items = []
    for tbl in doc.tables:
        for r_i, row in enumerate(tbl.rows):
            cells = [c.text.strip() for c in row.cells]
            if len(cells) >= 2:
                de, fr = cells[0], cells[1]
                if not de or not fr:
                    continue
                if r_i == 0 and ("de" in de.lower() and "fr" in fr.lower()):
                    continue
                items.append({"de": de, "fr": fr})
    for p in doc.paragraphs:
        t = p.text.strip()
        if ";" in t:
            parts = [s.strip() for s in t.split(";")]
            if len(parts) >= 2 and parts[0] and parts[1]:
                items.append({"de": parts[0], "fr": parts[1]})
    seen = set(); uniq = []
    for it in items:
        key = (normalize(it["de"]), normalize(it["fr"]))
        if key in seen:
            continue
        seen.add(key); uniq.append(it)
    return {"name": filename.rsplit(".",1)[0], "items": uniq}

def build_questions(pool, direction):
    out = []
    for it in pool:
        if direction == "DE→FR":
            out.append({"q": it["de"], "a": it["fr"]})
        else:
            out.append({"q": it["fr"], "a": it["de"]})
    return out

def all_possible_answers(pool, direction):
    if direction == "DE→FR":
        return list({it["fr"] for it in pool})
    return list({it["de"] for it in pool})

st.set_page_config(page_title=APP_TITLE, page_icon="🔤", layout="wide")
st.title(APP_TITLE)

st.session_state.setdefault("store", DEFAULT_STORE)
if "QS" not in st.session_state:
    st.session_state.QS = None  # dict or None

def start_quiz(pool, direction, mode, n):
    qa = build_questions(pool, direction)
    order = list(range(len(qa)))
    random.shuffle(order)
    order = order[:min(n, len(order))]
    st.session_state.QS = {
        "direction": direction,
        "mode": mode,
        "order": order,
        "i": 0,
        "score": 0,
        "qa": qa,
        "pool_answers": all_possible_answers(pool, direction),
        "history": []
    }

def current_q():
    QS = st.session_state.QS
    if QS is None: 
        return None
    if QS["i"] >= len(QS["order"]):
        return None
    idx = QS["order"][QS["i"]]
    return QS["qa"][idx]

def advance(user, ok, correct):
    QS = st.session_state.QS
    cur = current_q()
    if cur is None:
        return
    if ok:
        QS["score"] += 1
    QS["history"].append((cur["q"], user, ok, correct))
    QS["i"] += 1

tab_quiz, tab_manage = st.tabs(["🎯 Quiz", "📚 Sammlungen & Import"])

with tab_manage:
    st.subheader("Bestehende Sammlungen")
    for c in st.session_state["store"]["collections"]:
        st.markdown(f"- **{c['name']}** – {len(c['items'])} Einträge")
    st.divider()
    st.subheader("Import aus Word (.docx)")
    up = st.file_uploader("Word-Datei hochladen", type=["docx"])
    if up is not None:
        data = import_docx(up.read(), up.name)
        if data and len(data["items"])>0:
            if st.button(f"Sammlung '{data['name']}' importieren ({len(data['items'])} Einträge)"):
                cols = st.session_state["store"]["collections"]
                idx = next((i for i,c in enumerate(cols) if c['name']==data['name']), None)
                if idx is not None:
                    cols[idx] = data
                else:
                    cols.append(data)
                st.success(f"Importiert: {len(data['items'])} Einträge.")
    st.divider()
    st.subheader("Datenbank exportieren")
    js = json.dumps(st.session_state["store"], ensure_ascii=False, indent=2).encode("utf-8")
    st.download_button("vocab_store.json herunterladen", js, file_name="vocab_store.json", mime="application/json")

with tab_quiz:
    coll_opts = ["(alle)"] + [c["name"] for c in st.session_state["store"]["collections"]]
    sel_coll = st.selectbox("Sammlung", coll_opts, index=0)
    direction = st.radio("Richtung", ["DE→FR", "FR→DE"], horizontal=True)
    mode = st.radio("Quiztyp", ["Multiple Choice", "Freitext"], horizontal=True)
    n_q = st.slider("Anzahl Fragen", 5, 50, 10)

    pool = []
    for c in st.session_state["store"]["collections"]:
        if sel_coll != "(alle)" and c["name"] != sel_coll:
            continue
        pool.extend(c["items"])

    if st.button("Quiz starten", type="primary") and len(pool)>=4:
        start_quiz(pool, direction, mode, n_q)
        st.rerun()

    QS = st.session_state.QS
    if QS is None:
        st.info("Konfiguration wählen und **Quiz starten**.")
        st.stop()

    total = len(QS["order"])
    cur = current_q()
    if cur is None:
        st.success(f"Fertig! Punktzahl: {QS['score']}/{total}  ({round(100*QS['score']/max(1,total))}%)")
        st.dataframe(
            [{"Frage":h[0], "Ihre Antwort":h[1], "Korrekt":"Ja" if h[2] else "Nein", "Richtig":h[3]} for h in QS["history"]],
            use_container_width=True
        )
        c1, c2 = st.columns(2)
        if c1.button("Neues Quiz"):
            st.session_state.QS = None
            st.rerun()
        if c2.button("Nochmal gleiche Auswahl"):
            start_quiz(pool, QS["direction"], QS["mode"], len(QS["order"]))
            st.rerun()
        st.stop()

    st.info(f"Frage {QS['i']+1}/{total}  •  Punktzahl: {QS['score']}")
    st.write(f"**Übersetze:** {cur['q']}")

    with st.form(key=f"qform_{QS['i']}", clear_on_submit=True):
        if QS["mode"] == "Multiple Choice":
            distractors = [a for a in QS["pool_answers"] if normalize(a)!=normalize(cur["a"])]
            random.shuffle(distractors)
            opts = [cur["a"]] + distractors[:3]
            random.shuffle(opts)
            ans = st.radio("Wähle die richtige Übersetzung:", opts, index=None, key=f"radio_{QS['i']}")
        else:
            ans = st.text_input("Antwort eingeben", key=f"text_{QS['i']}")
        submitted = st.form_submit_button("Weiter", use_container_width=True)

    if submitted:
        if (QS["mode"] == "Multiple Choice" and ans is None) or (QS["mode"] == "Freitext" and not ans.strip()):
            st.warning("Bitte eine Antwort eingeben/auswählen.")
        else:
            ok = normalize(ans) == normalize(cur["a"])
            advance(ans, ok, cur["a"])
            st.rerun()
